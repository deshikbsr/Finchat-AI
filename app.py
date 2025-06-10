%%writefile app.py
import io
import json
import os
import sqlite3
import time
import bcrypt
import docx
from fpdf import FPDF
import google.generativeai as genai
from PyPDF2 import PdfReader
import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime, timedelta
from typing import List, Dict
from mistralai import Mistral
import requests
import pytesseract
from pdf2image import convert_from_bytes
import re
from openai import OpenAI
from PIL import Image
import pdfplumber
import torch
from torchvision import transforms
import uuid
import streamlit.components.v1 as components
from docx.shared import Pt
from streamlit_javascript import st_javascript


# ─── Configuration ──────────────────────────────────────────────────────────────
GOOGLE_API_KEY = os.environ.get(
    "GOOGLE_API_KEY",
    "AIzaSyANbVVzZACnYnus00xwwRRE01n34yoAmcU"  # fallback for dev/testing
)

MISTRAL_API_KEY = os.environ.get("MISTRAL_API_KEY", "DUW9f3t6nvZaNkEbxcrxYP4hLIrC3g7Y")
MISTRAL_ENDPOINT = "https://api.mistral.ai/v1/chat/completions"

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "sk-61f7f17d33bd4598b4dd61edd13af337")
DEEPSEEK_CLIENT = OpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com/v1"  # ensure /v1 is included
)

BRAND_COLORS = {
    "primary": "#2E86AB",
    "secondary": "#F18F01",
    "background": "#F7F7F7",
    "text": "#121111"
}

genai.configure(api_key=GOOGLE_API_KEY)
mistral_client = Mistral(api_key=MISTRAL_API_KEY)

# ─── Database Helpers ──────────────────────────────────────────────────────────
def init_db(db_path: str = "users.db"):
    conn = sqlite3.connect(db_path, check_same_thread=False)
    cursor = conn.cursor()
    # Create or migrate users table
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password BLOB NOT NULL,
            role TEXT NOT NULL,
            location_id TEXT,
            last_login TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    cursor.execute("PRAGMA table_info(users)")
    existing_cols = [col[1] for col in cursor.fetchall()]
    for col in ["location_id", "last_login", "created_at"]:
        if col not in existing_cols:
            try:
                cursor.execute(f"ALTER TABLE users ADD COLUMN {col} TEXT")
            except sqlite3.OperationalError:
                pass

    # Add subscriptions table
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS subscriptions (
            username TEXT PRIMARY KEY,
            lease_analysis BOOLEAN DEFAULT 0,
            deal_structuring BOOLEAN DEFAULT 0,
            offer_generator BOOLEAN DEFAULT 0,
            FOREIGN KEY(username) REFERENCES users(username)
        )
    """)

    # Interactions table
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS interactions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            feature TEXT,
            input_text TEXT,
            output_text TEXT,
            timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(username) REFERENCES users(username)
        )
        """
    )

    # Sessions table for persistent login
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS sessions (
            session_token TEXT PRIMARY KEY,
            username TEXT,
            expires_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(username) REFERENCES users(username)
        )
        """
    )
    conn.commit()
    return conn


def create_default_admin(conn):
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) FROM users")
    if cursor.fetchone()[0] == 0:
        admin_pwd = bcrypt.hashpw("admin123".encode(), bcrypt.gensalt())
        cursor.execute(
            "INSERT INTO users (username, password, role) VALUES (?, ?, ?)",
            ("admin", admin_pwd, "admin"),
        )
        conn.commit()


def verify_password(hashed: bytes, password: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed)


def create_session(conn, username: str, expires_days: int = 30) -> str:
    """Create a new session and return the session token."""
    session_token = str(uuid.uuid4())
    expires_at = datetime.now() + timedelta(days=expires_days)
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO sessions (session_token, username, expires_at) VALUES (?, ?, ?)",
        (session_token, username, expires_at)
    )
    conn.commit()
    return session_token


def validate_session(conn, session_token: str) -> Dict:
    """Validate a session token and return user info if valid."""
    cursor = conn.cursor()
    cursor.execute(
        "SELECT username, expires_at FROM sessions WHERE session_token = ?",
        (session_token,)
    )
    session = cursor.fetchone()
    if session and datetime.fromisoformat(session[1]) > datetime.now():
        cursor.execute(
            "SELECT username, role, location_id FROM users WHERE username = ?",
            (session[0],)
        )
        user = cursor.fetchone()
        if user:
            return {"username": user[0], "role": user[1], "location_id": user[2]}
    return None


def delete_session(conn, session_token: str):
    """Delete a session from the database."""
    cursor = conn.cursor()
    cursor.execute("DELETE FROM sessions WHERE session_token = ?", (session_token,))
    conn.commit()


def delete_user_sessions(conn, username: str):
    """Delete all sessions for a given user."""
    cursor = conn.cursor()
    cursor.execute("DELETE FROM sessions WHERE username = ?", (username,))
    conn.commit()


# ─── Authentication UI ─────────────────────────────────────────────────────────
def login_ui(conn):
    """Plain-style login UI with persistent session management."""
    # JavaScript to get/set session token in local storage
    components.html(
        """
        <script>
            function getSessionToken() {
                return localStorage.getItem('session_token') || '';
            }
            function setSessionToken(token) {
                localStorage.setItem('session_token', token);
                window.location.reload();
            }
            window.getSessionToken = getSessionToken;
            window.setSessionToken = setSessionToken;
        </script>
        """,
        height=0
    )

    # Check for existing session
    if "session_token" not in st.session_state:
        session_token = st_javascript("getSessionToken()")
        if session_token:
            user_info = validate_session(conn, session_token)
            if user_info:
                st.session_state.logged_in = True
                st.session_state.username = user_info["username"]
                st.session_state.role = user_info["role"]
                st.session_state.location_id = user_info.get("location_id")
                st.session_state.session_token = session_token
                st.rerun()

    if st.session_state.get("logged_in"):
        return

    # ─── Sidebar Styling ───────────────────────────────────────
    st.sidebar.title("🔑 Login / Register")
    st.sidebar.markdown(
        """
        <style>
            .sidebar .sidebar-content {
                background-color: #FFFFFF !important;
                color: #000000 !important;
                box-shadow: none !important;
            }
            .stTextInput>div>div>input,
            .stTextArea>div>div>textarea {
                background-color: #FFFFFF !important;
                color: #000000 !important;
                border: 1px solid #CCCCCC !important;
            }
            .stButton>button {
                background-color: #FFFFFF !important;
                color: #000000 !important;
                border: 1px solid #CCCCCC !important;
            }
        </style>
        """,
        unsafe_allow_html=True
    )

    login_tab, register_tab = st.sidebar.tabs(["Login", "Register"])

    # ─── LOGIN TAB ────────────────────────────────────────────────
    with login_tab:
        username = st.text_input("Username", key="login_username")
        password = st.text_input("Password", type="password", key="login_password")
        location_id = st.text_input("Login Key (use instead of user/pass)", key="login_location")

        if st.button("Log In", key="login_button"):
            # 1) Key-only login
            if location_id and not username and not password:
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT username, role FROM users WHERE location_id = ?",
                    (location_id,)
                )
                row = cursor.fetchone()
                if row:
                    session_token = create_session(conn, row[0])
                    st.session_state.logged_in = True
                    st.session_state.username = row[0]
                    st.session_state.role = row[1]
                    st.session_state.location_id = location_id
                    st.session_state.session_token = session_token
                    cursor.execute("UPDATE users SET last_login = CURRENT_TIMESTAMP WHERE username = ?", (row[0],))
                    conn.commit()
                    st_javascript(f"setSessionToken('{session_token}')")
                    st.rerun()
                else:
                    st.sidebar.error("❌ Invalid login key.")

            # 2) Fallback to username/password
            elif not username or not password:
                st.sidebar.error("Enter both username and password")

            else:
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT password, role, location_id FROM users WHERE username = ?",
                    (username,)
                )
                row = cursor.fetchone()

                if row and bcrypt.checkpw(password.encode(), row[0]):
                    session_token = create_session(conn, username)
                    st.session_state.logged_in = True
                    st.session_state.username = username
                    st.session_state.role = row[1]
                    st.session_state.location_id = row[2]
                    st.session_state.session_token = session_token
                    cursor.execute("UPDATE users SET last_login = CURRENT_TIMESTAMP WHERE username = ?", (username,))
                    conn.commit()
                    st_javascript(f"setSessionToken('{session_token}')")
                    st.rerun()
                else:
                    st.sidebar.error("Invalid username or password")
                    time.sleep(1)

    # ─── REGISTER TAB ────────────────────────────────────────────
    with register_tab:
        new_user = st.text_input("New Username", key="reg_username")
        new_pass = st.text_input("New Password", type="password", key="reg_password")
        confirm_pass = st.text_input("Confirm Password", type="password", key="reg_confirm")
        user_role = "user"

        if st.button("Create User", key="reg_button"):
            if not new_user or not new_pass:
                st.error("Username and password are required")
            elif new_pass != confirm_pass:
                st.error("Passwords do not match")
            elif len(new_pass) < 8:
                st.error("Password must be at least 8 characters")
            else:
                try:
                    location_key = str(uuid.uuid4())
                    hashed = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt())
                    cursor = conn.cursor()
                    cursor.execute(
                        "INSERT INTO users (username, password, role, location_id) VALUES (?, ?, ?, ?)",
                        (new_user, hashed, user_role, location_key)
                    )
                    session_token = create_session(conn, new_user)
                    conn.commit()
                    st.success(f"User '{new_user}' created successfully.")
                    st.info(f"🔑 **Your login key** (save this!): `{location_key}`")
                    st.session_state.logged_in = True
                    st.session_state.username = new_user
                    st.session_state.role = user_role
                    st.session_state.location_id = location_key
                    st.session_state.session_token = session_token
                    st_javascript(f"setSessionToken('{session_token}')")
                    time.sleep(1)
                    st.rerun()
                except sqlite3.IntegrityError:
                    st.error("That username already exists")

    # ─── MAIN PANE WELCOME BANNER ───────────────────────────────
    if not st.session_state.get("logged_in"):
        st.markdown(
            """
            <div style="
                height: 80vh;
                display: flex;
                flex-direction: column;
                align-items: center;
                justify-content: center;
                text-align: center;
            ">
                <h1 style="color: #2E86AB; font-weight: normal; margin-bottom: 0.2em;">
                    Welcome to Finchat AI Bot
                </h1>
                <p style="color: #555555; font-size: 1.1em; margin-top: 0;">
                    🤖 Powered by Alphax — crafting real estate insights in seconds!
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )
        return


# ─── AI Helper Functions ───────────────────────────────────────────────────────
def call_gemini(
    feature: str,
    content: str,
    temperature: float = 0.7
) -> str:
    """Call Google's Gemini model with proper temperature handling."""
    system_prompts = {
        "lease_analysis": (
            "You are a real estate document expert. Analyze the provided lease agreement "
            "and provide a comprehensive summary, including key terms and potential risks."
        ),
        "deal_strategy": (
            "You are a creative real estate strategist. Based on the provided deal details, "
            "suggest structuring options with pros, cons, and negotiation tactics."
        ),
        "offer_generator": (
            "You are a real estate transaction specialist. Generate a professional purchase offer "
            "with all essential clauses formatted for the jurisdiction."
        ),
        "chatbot": (
            "You are a knowledgeable assistant that answers questions based on the user's past interactions."
        )
    }

    try:
        import google.generativeai as genai
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"SYSTEM: {system_prompts.get(feature, '')}\n\nUSER: {content}"
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=temperature,
                top_p=0.95,
                top_k=40,
                max_output_tokens=8192
            )
        )
        return response.text
    except Exception as e:
        st.error(f"Gemini API error: {e}")
        return f"Error: {e}"


def call_mistral(
    messages: List[Dict[str, str]],
    temperature: float = 0.2,
    top_p: float = 1.0,
    max_tokens: int = 1024,
    stop: List[str] = None,
    stream: bool = False,
    user: str = None,
    logit_bias: Dict[int, float] = None,
) -> str:
    """
    Call Mistral API with extended parameter support.
    """
    payload = {
        "model": "mistral-small-latest",
        "messages": messages,
        "temperature": temperature,
        "top_p": top_p,
        "max_tokens": max_tokens,
    }
    if stop:
        payload["stop"] = stop
    if user:
        payload["user"] = user
    if logit_bias:
        payload["logit_bias"] = logit_bias

    headers = {"Authorization": f"Bearer {MISTRAL_API_KEY}", "Content-Type": "application/json"}
    resp = requests.post(MISTRAL_ENDPOINT, json=payload, headers=headers, stream=stream)
    resp.raise_for_status()
    data = resp.json()

    if stream:
        return "".join(chunk.get("content", "") for chunk in data.get("choices", []))
    return data["choices"][0]["message"]["content"]


def call_deepseek(
    messages: List[Dict[str, str]],
    model: str = "deepseek-chat",
    temperature: float = 0.7,
    max_tokens: int = 2000,
    top_p: float = 1.0,
    frequency_penalty: float = 0.0,
    presence_penalty: float = 0.0,
    stream: bool = False,
) -> str:
    """
    Call DeepSeek via the OpenAI-compatible SDK.
    """
    try:
        resp = DEEPSEEK_CLIENT.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=top_p,
            frequency_penalty=frequency_penalty,
            presence_penalty=presence_penalty,
            stream=stream
        )
        if stream:
            return "".join(chunk.choices[0].delta.content for chunk in resp)
        return resp.choices[0].message.content
    except Exception as e:
        return f"Error processing request with DeepSeek: {str(e)}"


def save_interaction(conn, feature: str, input_text: str, output_text: str):
    if st.session_state.get("username"):
        conn.execute(
            "INSERT INTO interactions (username, feature, input_text, output_text) VALUES (?, ?, ?, ?)",
            (st.session_state.username, feature, input_text, output_text),
        )
        conn.commit()


# Helper to strip Markdown tokens

def strip_markdown(text: str) -> str:
    # remove bold/italic markers
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*",   r"\1", text)
    # remove ATX headings
    text = re.sub(r"^\s*#{1,6}\s*", '', text, flags=re.MULTILINE)
    # remove list bullets
    text = re.sub(r"^\s*[-*+]\s*",   '', text, flags=re.MULTILINE)
    return text.strip()


def lease_summarization_ui(conn):
    """Lease Summarization: upload PDF and get either full-document or page-by-page summaries with model selection, persisting results for chatbot usage, and answering questions from an uploaded document."""
    st.header("📄 Lease Summary")

    # Clear previous summary state
    if st.button("Clear Summary", key="clear_lease_summary"):
        for k in ["last_file", "last_summary", "last_mode", "last_engine", "last_questions", "last_answers"]:
            st.session_state.pop(k, None)
        st.success("Cleared previous summary and questions.")
        st.rerun()

    st.markdown(
        "Upload your lease PDF and receive a concise summary—choose to process the entire document at once or summarize each page individually."
    )

    uploaded_file = st.file_uploader(
        "Upload Lease Document (PDF)", type=["pdf"], key="lease_file_uploader"
    )
    if 'last_file' in st.session_state and uploaded_file:
        if st.session_state.last_file != uploaded_file.name:
            for k in ['last_summary', 'last_mode', 'last_engine', 'last_questions', 'last_answers']:
                st.session_state.pop(k, None)
    if not uploaded_file:
        return

    ai_engine = st.radio(
        "Select AI Model",
        ["In depth summarisation", "General Summary", "General Summary Pro"],
        index=0,
        horizontal=True,
        key="lease_ai_engine"
    )
    summary_mode = st.radio(
        "Summary Mode",
        ["Full Document", "Page-by-Page"],
        index=1,
        horizontal=True,
        key="lease_summary_mode"
    )

    # --- New: Initialize question-related session state ---
    if 'last_questions' not in st.session_state:
        st.session_state.last_questions = []
        st.session_state.last_answers = []

    # Display existing summary if available
    if 'last_summary' in st.session_state and st.session_state['last_file'] == uploaded_file.name:
        mode = st.session_state['last_mode']
        engine = st.session_state['last_engine']
        raw = st.session_state['last_summary']

        # Display summary
        if mode == 'Full Document':
            summary_content = raw
            st.subheader(f"Full Document Summary ({engine})")
            st.write(summary_content)
        else:
            parts = raw
            st.subheader(f"Page-by-Page Summary ({engine})")
            for part in parts:
                st.write(part)
            summary_content = "\n\n".join(parts)

        st.divider()
        st.markdown("### 📥 Export Styled Summary")

        file_base = uploaded_file.name.rsplit(".", 1)[0]
        file_name = st.text_input("Filename (no extension):", value=file_base, key="lease_export_name")

        # Clean and ensure Latin-1 encoding for PDF
        paragraphs_pdf = [
            strip_markdown(p)
            .encode('latin-1', 'replace')
            .decode('latin-1')
            for p in summary_content.split("\n\n") if p.strip()
        ]
        paragraphs_word = [strip_markdown(p) for p in summary_content.split("\n\n") if p.strip()]

        class PDF(FPDF):
            def header(self):
                self.set_font('Helvetica', 'B', 16)
                self.set_text_color(44, 134, 171)
                self.cell(0, 10, 'Lease Summary Report', ln=True, align='C')
                self.set_font('Helvetica', 'I', 10)
                self.set_text_color(85, 85, 85)
                self.cell(0, 6, f'Mode: {mode} | Engine: {engine} | Generated: {datetime.now():%B %d, %Y}', ln=True, align='C')
                self.ln(8)
                self.set_draw_color(200, 200, 200)
                self.line(10, self.get_y(), 200, self.get_y())
                self.ln(5)

            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 8)
                self.set_text_color(85, 85, 85)
                self.cell(0, 10, f'Property Deals AI | Page {self.page_no()}/{{nb}}', align='C')

            def add_section(self, title, content, page_number=None):
                if title:
                    self.set_font('Helvetica', 'B', 12)
                    self.set_text_color(44, 134, 171)
                    self.cell(0, 8, title, ln=True)
                self.set_font('Helvetica', '', 11)
                self.set_text_color(0, 0, 0)
                if isinstance(content, list):
                    for item in content:
                        self.multi_cell(0, 6, item)
                        self.ln(1)
                else:
                    self.multi_cell(0, 6, content)
                self.ln(4)

        pdf = PDF()
        pdf.alias_nb_pages()
        pdf.set_left_margin(15)
        pdf.set_right_margin(15)
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        if mode == 'Full Document':
            pdf.add_section('Full Document Summary', paragraphs_pdf)
        else:
            for para in paragraphs_pdf:
                pdf.add_section('', [para])

        # --- New: Add questions and answers to PDF if available ---
        if st.session_state.last_questions and st.session_state.last_answers:
            pdf.add_page()
            pdf.add_section('Questions and Answers', [
                f"Q{i+1}: {q}\nA{i+1}: {a}"
                for i, (q, a) in enumerate(zip(st.session_state.last_questions, st.session_state.last_answers))
            ])

        pdf_bytes = pdf.output(dest='S').encode('latin-1', 'replace')
        st.download_button(
            "Download Styled PDF",
            data=pdf_bytes,
            file_name=f"{file_name}.pdf",
            mime="application/pdf",
            key="lease_export_pdf"
        )

        # Word export
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        doc.add_heading('Lease Summary', level=1)
        doc.add_paragraph(f"Mode: {mode} | Engine: {engine}")
        doc.add_paragraph("")
        for para in paragraphs_word:
            doc.add_paragraph(para)

        # --- New: Add questions and answers to Word document ---
        if st.session_state.last_questions and st.session_state.last_answers:
            doc.add_heading('Questions and Answers', level=2)
            for i, (q, a) in enumerate(zip(st.session_state.last_questions, st.session_state.last_answers)):
                doc.add_paragraph(f"Q{i+1}: {q}", style='Normal')
                doc.add_paragraph(f"A{i+1}: {a}", style='Normal')
                doc.add_paragraph("")

        buf = io.BytesIO()
        doc.save(buf)
        st.download_button(
            "Download Styled Word",
            data=buf.getvalue(),
            file_name=f"{file_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="lease_export_word"
        )

        # --- New: Question Document Upload Section ---
        st.divider()
        st.markdown("### ❓ Ask Questions from Document")
        st.markdown("Upload a document (PDF, DOCX, or TXT) containing questions about the lease summary.")

        question_file = st.file_uploader(
            "Upload Question Document (PDF, DOCX, TXT)",
            type=["pdf", "docx", "txt"],
            key="question_file_uploader"
        )

        if question_file:
            try:
                # Extract text from the uploaded question document
                if question_file.type == "application/pdf":
                    reader = PdfReader(question_file)
                    question_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                elif question_file.type == "text/plain":
                    question_text = question_file.read().decode("utf-8")
                elif question_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = docx.Document(question_file)
                    question_text = "\n".join(p.text for p in doc.paragraphs)
                else:
                    st.error("Unsupported file type.")
                    return

                # Extract questions (assuming questions are separated by newlines or numbered)
                questions = [
                    q.strip() for q in question_text.split("\n")
                    if q.strip() and (q.strip().endswith("?") or re.match(r"^\d+\.|\d+\)", q))
                ]
                if not questions:
                    st.warning("No questions detected in the document. Ensure questions end with '?' or are numbered.")
                    return

                st.session_state.last_questions = questions
                st.session_state.last_answers = []

                # Generate answers using the lease summary as context
                with st.spinner("Generating answers to your questions..."):
                    for i, question in enumerate(questions):
                        prompt = (
                            f"Based on the following lease summary, provide a detailed explanation or reason for this question:\n\n"
                            f"Summary:\n{summary_content}\n\n"
                            f"Question: {question}"
                        )
                        answer = call_deepseek(
                            messages=[
                                {"role": "system", "content": "You are a real estate expert providing explanations based on a lease summary."},
                                {"role": "user", "content": prompt}
                            ],
                            model="deepseek-chat",
                            temperature=0.3,
                            max_tokens=512
                        )
                        st.session_state.last_answers.append(answer)
                        save_interaction(conn, "lease_question_answer", question, answer)

                # Display questions and answers
                st.subheader("Questions and Answers")
                for i, (q, a) in enumerate(zip(st.session_state.last_questions, st.session_state.last_answers)):
                    with st.expander(f"Q{i+1}: {q}"):
                        st.markdown(f"**A{i+1}:** {a}")

            except Exception as e:
                st.error(f"Failed to process question document: {e}")

    if st.button("Generate Summary", key="lease_generate_button"):
        try:
            reader = PdfReader(uploaded_file)
            pages = [page.extract_text() or "" for page in reader.pages]
        except Exception:
            st.error("Failed to extract text from the PDF.")
            return

        if not any(pages):
            st.error("No readable text found in the PDF.")
            return

        st.session_state['last_file'] = uploaded_file.name
        st.session_state['last_mode'] = summary_mode
        st.session_state['last_engine'] = ai_engine

        if summary_mode == "Full Document":
            text = "\n".join(pages)
            with st.spinner("Summarizing full document..."):
                summaries = []
                chunks = [text[i:i+15000] for i in range(0, len(text), 15000)] if len(text) > 15000 else [text]
                for chunk in chunks:
                    prompt = (
                        "Summarize this portion of the lease agreement in clear, concise language, "
                        "preserving all key details:\n\n" + chunk
                    )
                    summaries.append(call_deepseek(messages=[{"role":"user","content":prompt}], model="deepseek-chat", temperature=0.3, max_tokens=1024))
                final = "\n\n".join(summaries)
            st.subheader("Full Document Summary")
            st.write(final)
            save_interaction(conn, "lease_summary_full", uploaded_file.name, final)
            st.session_state['last_summary'] = final
        else:
            parts = []
            st.subheader("Page-by-Page Summaries")
            for i, pg in enumerate(pages, start=1):
                if not pg.strip():
                    parts.append("(no text detected)")
                else:
                    with st.spinner(f"Summarizing page {i}..."):
                        prompt = (
                            f"Summarize page {i} of this lease agreement in clear, concise language, covering all information:\n\n{pg}"
                        )
                        summary = call_deepseek(messages=[{"role":"user","content":prompt}], model="deepseek-chat", temperature=0.3, max_tokens=512)
                        parts.append(summary)
            save_interaction(conn, "lease_summary_pagewise", uploaded_file.name, json.dumps({f"page_{i}": pages[i-1] for i in range(1, len(pages)+1)}))
            st.session_state['last_summary'] = parts
        st.rerun()

def deal_structuring_ui(conn):
    """Enhanced deal structuring with persistent strategy chat until cleared."""
    st.header("💡 Creative Deal Structuring Bot")
    st.markdown("Get AI-powered strategies for your property deals")

    if "deal_strategy_memory" not in st.session_state:
        st.session_state.deal_strategy_memory = []
        st.session_state.last_strategies = None
        st.session_state.strategy_confidences = {}

    if st.button("Clear Strategies", key="clear_strategies"):
        st.session_state.deal_strategy_memory.clear()
        st.session_state.last_strategies = None
        st.session_state.strategy_confidences = {}
        st.rerun()

    for role, msg in st.session_state.deal_strategy_memory:
        st.chat_message(role).write(msg)

    with st.expander("Deal Details", expanded=True):
        property_type = st.selectbox("Property Type", ["Residential", "Commercial", "Mixed-Use", "Land"])
        deal_stage = st.selectbox("Deal Stage", ["Pre-offer", "Under Contract", "Rehab Planning", "Exit Strategy"])
        financials = st.text_area("Financial Parameters")
        market_conditions = st.text_area("Market Conditions")
        special_considerations = st.text_area("Special Considerations")

    with st.expander("Strategy Preferences"):
        col1, col2 = st.columns(2)
        with col1:
            risk_tolerance = st.select_slider("Risk Tolerance", ["Conservative", "Moderate", "Aggressive"])
            creativity_level = st.select_slider("Creativity Level", ["Standard", "Creative", "Outside-the-box"])
        with col2:
            timeframe = st.selectbox(
                "Investment Horizon",
                ["Short-term (0-2 years)", "Medium-term (2-5 years)", "Long-term (5+ years)"]
            )
            capital_available = st.selectbox("Capital Availability", ["Limited", "Moderate", "Substantial"])

    ai_model = st.radio("AI Model", ["Gemini", "Mistral", "DeepSeek"], horizontal=True)

    if st.button("Generate Strategies", type="primary", key="gen_strat"):
        prompt = (
            f"Property Type: {property_type}\n"
            f"Deal Stage: {deal_stage}\n"
            f"Financial Parameters: {financials}\n"
            f"Market Conditions: {market_conditions}\n"
            f"Special Considerations: {special_considerations}\n\n"
            f"Generate {risk_tolerance.lower()} strategies with {creativity_level.lower()} approaches "
            f"for a {timeframe} investment using {capital_available.lower()} capital."
        )
        with st.spinner("Developing strategies..."):
            if ai_model == "Gemini":
                strategies = call_gemini("deal_strategy", prompt)
            elif ai_model == "Mistral":
                messages = [
                    {"role": "system", "content": "You are a real estate investment strategist. Provide creative deal structuring options."},
                    {"role": "user", "content": prompt}
                ]
                strategies = call_mistral(messages=messages)
            else:
                messages = [
                    {"role": "system", "content": "You are an expert real estate strategist. Suggest creative deal structures with pros/cons."},
                    {"role": "user", "content": prompt}
                ]
                strategies = call_deepseek(messages)

        st.session_state.deal_strategy_memory.append(("assistant", strategies))
        st.session_state.last_strategies = strategies
        st.chat_message("assistant").write(strategies)
        st.subheader("Recommended Strategies")
        st.markdown(strategies)

        matches = re.findall(
            r"Strategy\s+(\d+):\s*(.*?)(?=(?:Strategy\s+\d+:)|\Z)",
            strategies,
            flags=re.S
        )
        if matches:
            for num, _ in matches:
                strategy_key = f"Strategy {num}"
                if strategy_key not in st.session_state.strategy_confidences:
                    st.session_state.strategy_confidences[strategy_key] = 7
        else:
            if "Strategy 1" not in st.session_state.strategy_confidences:
                st.session_state.strategy_confidences["Strategy 1"] = 7

    strategies = st.session_state.get("last_strategies")
    if strategies:
        matches = re.findall(
            r"Strategy\s+(\d+):\s*(.*?)(?=(?:Strategy\s+\d+:)|\Z)",
            strategies,
            flags=re.S
        )
        if matches:
            strategy_dict = {f"Strategy {num}": text.strip() for num, text in matches}
        else:
            strategy_dict = {"Strategy 1": strategies.strip()}

        labels = list(strategy_dict.keys())
        selected_label = st.selectbox("Which strategy do you prefer?", labels, key="eval_choice")
        selected_text = strategy_dict[selected_label]

        st.markdown(f"**{selected_label}**")
        st.markdown(selected_text)

        confidence = st.slider(
            "Confidence in this strategy",
            1, 10,
            value=st.session_state.strategy_confidences.get(selected_label, 7),
            key=f"conf_{selected_label.replace(' ', '_')}"
        )

        st.session_state.strategy_confidences[selected_label] = confidence

        if st.button("Refine Strategy", key="refine_strat"):
            feedback = f"{selected_label} with confidence {confidence}/10"
            st.session_state.deal_strategy_memory.append(("user", feedback))

            refinement_prompt = (
                f"Refine this single strategy based on user feedback:\n\n"
                f"{selected_text}\n\n"
                f"Feedback: {feedback}"
            )
            if ai_model == "Gemini":
                refinement = call_gemini("deal_strategy", refinement_prompt)
            elif ai_model == "Mistral":
                messages = [
                    {"role": "system", "content": "Refine the selected strategy based on user feedback."},
                    {"role": "user", "content": refinement_prompt}
                ]
                refinement = call_mistral(messages=messages)
            else:
                messages = [
                    {"role": "system", "content": "Refine this real estate strategy based on the provided feedback."},
                    {"role": "user", "content": refinement_prompt}
                ]
                refinement = call_deepseek(messages)

            st.session_state.deal_strategy_memory.append(("assistant", refinement))
            st.chat_message("assistant").write(refinement)
            save_interaction(conn, "deal_strategy_refinement", selected_text, refinement)


def build_guided_prompt(details: dict, detail_level: str) -> str:
    """
    Construct a detailed prompt from guided form data to generate a real estate purchase agreement.
    """
    buyer = details['parties']['buyer']
    buyer_rep = details['parties'].get('buyer_rep', '')
    seller = details['parties']['seller']
    seller_rep = details['parties'].get('seller_rep', '')
    buyer_line = f"- Buyer: {buyer}{f' (Represented by: {buyer_rep})' if buyer_rep else ''}"
    seller_line = f"- Seller: {seller}{f' (Represented by: {seller_rep})' if seller_rep else ''}"

    address = details['property']['address']
    county = details['property'].get('county', '')
    address_line = f"- Property Address: {address}"
    county_line = f"- County: {county}" if county else ''

    price = details['financial']['price_fmt']
    earnest = details['financial']['earnest_fmt']
    price_line = f"- Purchase Price: {price}"
    earnest_line = f"- Earnest Money Deposit: {earnest}"

    closing = details['dates']['closing']
    expiry = details['dates']['expiry']
    closing_line = f"- Proposed Closing Date: {closing}"
    expiry_line = f"- Offer Expiration: {expiry} hours from signing"

    financing = details['terms'].get('financing', '')
    contingencies = details['terms'].get('contingencies', [])
    contingencies_str = ', '.join(contingencies) if contingencies else 'None'
    special_terms = details['terms'].get('special', '')
    financing_line = f"- Financing Type: {financing}"
    contingencies_line = f"- Contingencies: {contingencies_str}"
    special_line = f"- Special Terms: {special_terms}" if special_terms else ''

    jurisdiction = details['terms'].get('jurisdiction', '')
    jurisdiction_line = f"- Governing Law: {jurisdiction}" if jurisdiction else ''

    sections = [
        "Generate a professional real estate purchase agreement with the following details:",
        buyer_line, seller_line, address_line, county_line,
        price_line, earnest_line, closing_line, expiry_line,
        financing_line, contingencies_line, special_line, jurisdiction_line,
        f"Level of Detail: {detail_level}."
    ]
    return "\n".join([s for s in sections if s])


def offer_generator_ui(conn):
    st.header("✍️ Advanced Offer Generator")
    st.markdown(
        """
        <style>
        .offer-section { background-color: #f0f2f6; border-radius: 10px; padding: 15px; margin-bottom: 20px; }
        .offer-highlight { background-color: #fffacd; padding: 2px 5px; border-radius: 3px; }
        </style>
        """, unsafe_allow_html=True
    )

    if 'offer_stage' not in st.session_state:
        st.session_state.update({
            'offer_stage': 'input_method',
            'offer_data': {},
            'generated_offer': None,
            'edited_offer': None,
            'review_comments': []
        })

    stages = ["input_method", "details_entry", "offer_generation", "review_edit", "export"]
    labels = ["Input Method", "Details Entry", "Offer Generation", "Review & Edit", "Export"]
    idx = stages.index(st.session_state.offer_stage)
    cols = st.columns(len(stages))
    for i, label in enumerate(labels):
        with cols[i]:
            if i < idx:
                st.success(f"✓ {label}")
            elif i == idx:
                st.info(f"→ {label}")
            else:
                st.caption(label)

    if st.session_state.offer_stage == 'input_method':
        st.markdown("### 1. Select Input Method")
        method = st.radio(
            "How would you like to create your offer?",
            ["Guided Form", "Free Text", "Upload Existing", "Template Library"],
            horizontal=True,
            key="offer_input_method"
        )
        st.session_state.offer_data['input_method'] = method
        with st.expander("AI Configuration"):
            ai_model = st.radio(
                "AI Model Preference", ["Gemini", "Mistral", "DeepSeek"], horizontal=True, key="offer_ai_model"
            )
            creativity = st.slider("Creativity Level", 0.0, 1.0, 0.3, key="offer_creativity")
            detail_level = st.select_slider(
                "Detail Level", options=["Minimal","Standard","Comprehensive"],
                value="Standard", key="offer_detail_level"
            )
            st.session_state.offer_data.update({
                'ai_model': ai_model,
                'creativity': creativity,
                'detail_level': detail_level
            })
        if st.button("Continue to Details", key="btn_continue_details"):
            st.session_state.offer_stage = 'details_entry'
            st.rerun()

    elif st.session_state.offer_stage == 'details_entry':
        st.markdown("### 2. Enter Offer Details")
        method = st.session_state.offer_data['input_method']
        if method == 'Guided Form':
            with st.form("offer_details_form"):
                st.markdown('<div class="offer-section">', unsafe_allow_html=True)
                st.markdown('#### Basic Information')
                c1, c2, c3 = st.columns(3)
                with c1:
                    buyer = st.text_input("Buyer Full Name*", key="offer_buyer")
                    buyer_rep = st.text_input("Buyer\'s Representative", key="offer_buyer_rep")
                with c2:
                    seller = st.text_input("Seller Full Name*", key="offer_seller")
                    seller_rep = st.text_input("Seller\'s Representative", key="offer_seller_rep")
                with c3:
                    address = st.text_input("Property Address*", key="offer_address")
                    county = st.text_input("County", key="offer_county")
                st.markdown('#### Financial Terms')
                c1, c2, c3 = st.columns(3)
                with c1:
                    price = st.number_input("Purchase Price*", min_value=1000, step=1000, key="offer_price")
                with c2:
                    earnest = st.number_input("Earnest Money Deposit*", min_value=0, step=1000, key="offer_earnest")
                with c3:
                    closing = st.date_input("Proposed Closing Date*", min_value=datetime.now().date(), key="offer_closing")
                st.markdown('#### Terms & Conditions')
                c1, c2 = st.columns(2)
                with c1:
                    financing = st.selectbox("Financing Type*", ["Cash","Conventional Loan","FHA","VA","Seller Financing","Other"], key="offer_financing")
                    if financing == "Other": st.text_input("Specify Financing Type", key="offer_financing_other")
                with c2:
                    cont = st.multiselect("Contingencies", ["Inspection","Appraisal","Financing","Title Review","HOA Approval","Other"], key="offer_contingencies")
                    if "Other" in cont: st.text_input("Specify Other Contingency", key="offer_contingencies_other")
                st.markdown('#### Additional Provisions')
                terms = st.text_area("Special Terms/Conditions", key="offer_special_terms")
                st.markdown('#### Jurisdiction & Expiry')
                c1, c2 = st.columns(2)
                with c1:
                    jurisdiction = st.selectbox("Governing Law", ["State Default","California","Texas","New York","Florida","Other"], key="offer_jurisdiction")
                with c2:
                    expiry = st.number_input("Offer Expiration (hours)", min_value=1, max_value=168, value=48, key="offer_expiry")
                st.markdown('</div>', unsafe_allow_html=True)

                submitted = st.form_submit_button("Generate Offer Draft")
                if submitted:
                    missing = []
                    for field, msg in {
                        "offer_buyer": "Buyer required",
                        "offer_seller": "Seller required",
                        "offer_address": "Address required",
                        "offer_price": "Price >=1000",
                        "offer_earnest": "Earnest required",
                        "offer_closing": "Closing date required"
                    }.items():
                        if not st.session_state.get(field):
                            missing.append(msg)
                    if missing:
                        for m in missing:
                            st.error(m)
                    else:
                        st.session_state.offer_data['details'] = {
                            'parties': {
                                'buyer': buyer,
                                'buyer_rep': buyer_rep,
                                'seller': seller,
                                'seller_rep': seller_rep
                            },
                            'property': {
                                'address': address,
                                'county': county
                            },
                            'financial': {
                                'price': price,
                                'earnest': earnest,
                                'price_fmt': f"${price:,}",
                                'earnest_fmt': f"${earnest:,}"
                            },
                            'dates': {
                                'closing': closing.strftime("%B %d, %Y"),
                                'expiry': expiry
                            },
                            'terms': {
                                'financing': financing,
                                'contingencies': cont,
                                'special': terms,
                                'jurisdiction': jurisdiction
                            }
                        }
                        st.session_state.offer_stage = 'offer_generation'
                        st.rerun()

        elif method == 'Free Text':
            st.markdown("Enter deal details (min 50 chars):")
            text = st.text_area("Deal Details", key="offer_free_text", height=200)
            if st.button("Generate Offer Draft Free Text", key="btn_ft_draft"):
                if len(text) < 50:
                    st.error("Please add more detail.")
                else:
                    st.session_state.offer_data['details'] = {'free_text': text}
                    st.session_state.offer_stage = 'offer_generation'
                    st.rerun()

        elif method == 'Upload Existing':
            uploaded = st.file_uploader("Upload Document", type=["pdf","docx","txt"], key="offer_upload")
            if uploaded and st.button("Analyze & Improve Upload", key="btn_upload_analyze"):
                if uploaded.type == "application/pdf":
                    reader = PdfReader(uploaded)
                    doc_text = "\n".join(p.extract_text() or "" for p in reader.pages)
                elif uploaded.type == "text/plain":
                    doc_text = uploaded.read().decode("utf-8")
                else:
                    doc_text = "\n".join(p.text for p in doc.paragraphs)
                st.session_state.offer_data['details'] = {'uploaded': doc_text}
                st.session_state.offer_stage = 'offer_generation'
                st.rerun()

        else:
            st.markdown("### Template Library")
            templates = {
                "Residential": "templates/standard_residential.json",
                "Commercial": "templates/commercial_lease_purchase.json",
                "Seller Financing": "templates/seller_financing.json",
                "1031 Exchange": "templates/1031_exchange.json"
            }
            choice = st.selectbox("Select Template", list(templates.keys()), key="offer_template")
            with st.expander("Preview"):
                try:
                    st.json(json.load(open(templates[choice])))
                except Exception:
                    st.warning("Preview unavailable")
            if st.button("Use Template", key="btn_use_template"):
                try:
                    data = json.load(open(templates[choice]))
                    st.session_state.offer_data['details'] = data
                    st.session_state.offer_stage = 'offer_generation'
                    st.rerun()
                except Exception:
                    st.error("Failed to load")

        if st.button("← Back", key="btn_back_to_input"):
            st.session_state.offer_stage = 'input_method'
            st.rerun()

    if st.session_state.offer_stage == 'offer_generation':
        d = st.session_state.offer_data
        if d['input_method'] == 'Guided Form':
            prompt = build_guided_prompt(d['details'], d['detail_level'])
        elif d['input_method'] == 'Free Text':
            prompt = f"Draft a purchase agreement:\n\n{d['details']['free_text']}"
        elif d['input_method'] == 'Upload Existing':
            prompt = f"Improve this draft:\n\n{d['details']['uploaded']}"
        else:
            prompt = f"Generate from template:\n\n{json.dumps(d['details'], indent=2)}"
        prompt += f"\n\nDetail Level: {d['detail_level']}"

        with st.spinner("Generating..."):
            if d['ai_model'] == 'Gemini':
                offer = call_gemini('offer_generator', prompt)
            elif d['ai_model'] == 'Mistral':
                messages = [
                    {'role': 'system', 'content': 'You are a real estate attorney.'},
                    {'role': 'user', 'content': prompt}
                ]
                offer = call_mistral(messages, temperature=d['creativity'])
            else:
                messages = [
                    {'role': 'system', 'content': 'You are a legal expert drafting a real estate purchase agreement.'},
                    {'role': 'user', 'content': prompt}
                ]
                offer = call_deepseek(messages, temperature=d['creativity'])

            st.session_state.generated_offer = offer
            save_interaction(conn, 'offer_generator', prompt, offer)

        st.subheader("Generated Offer")
        st.markdown(offer, unsafe_allow_html=True)
        if st.button("Proceed to Review"): st.session_state.offer_stage = 'review_edit'; st.rerun()
        if st.button("← Back"): st.session_state.offer_stage = 'details_entry'; st.rerun()

    if st.session_state.offer_stage == 'review_edit':
        edited = st.text_area(
            "Edit draft", value=st.session_state.generated_offer,
            height=300, key='offer_edit'
        )
        if edited != st.session_state.edited_offer:
            st.session_state.edited_offer = edited

        st.markdown("#### Comments")
        new_c = st.text_input("Add comment", key='offer_new_comment')
        if st.button("Add Comment") and new_c:
            ts = datetime.now().strftime("%Y-%m-%d %H:%M")
            st.session_state.review_comments.append({
                'ts': ts, 'text': new_c, 'resolved': False
            })
            st.rerun()

        for i, c in enumerate(st.session_state.review_comments):
            cols = st.columns([1, 8, 1])
            with cols[0]:
                st.markdown(f"**{c['ts']}**")
            with cols[1]:
                st.markdown(f"{'✓' if c['resolved'] else '◯'} {c['text']}")
            with cols[2]:
                if not c['resolved'] and st.button('Resolve', key=f'res_{i}'):
                    c['resolved'] = True
                    st.rerun()

        if st.button('← Back'):
            st.session_state.offer_stage = 'offer_generation'
            st.rerun()
        if st.button('Proceed to Export'):
            st.session_state.offer_stage = 'export'
            st.rerun()

    if st.session_state.offer_stage == 'export':
        content = st.session_state.edited_offer or st.session_state.generated_offer
        if st.checkbox('Include Comments', value=True):
            content += (
                "\n\n---\n## Comments\n" +
                "\n".join([f"- [{c['ts']}] {c['text']}" for c in st.session_state.review_comments])
            )

        fmt = st.selectbox('Format', ['PDF', 'Word', 'Text', 'HTML'], key='offer_export_format')
        name = st.text_input('File Name', 'property_offer', key='offer_export_name')

        if st.button('Download'):
            if fmt == 'PDF':
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font('Arial', size=12)
                for line in content.split('\n'):
                    pdf.multi_cell(0, 6, line.encode('latin-1', 'replace').decode('latin-1'))
                st.download_button(
                    'Download PDF', pdf.output(dest='S').encode('latin-1'),
                    f"{name}.pdf", "application/pdf"
                )
            elif fmt == 'Word':
                doc = docx.Document()
                for line in content.split('\n'):
                    doc.add_paragraph(line)
                buf = io.BytesIO()
                doc.save(buf)
                st.download_button(
                    'Download Word', buf.getvalue(),
                    f"{name}.docx",
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
            elif fmt == 'HTML':
                html = f"<pre>{content}</pre>"
                st.download_button(
                    'Download HTML', html.encode(),
                    f"{name}.html", 'text/html'
                )
            else:
                st.download_button(
                    'Download Text', content.encode(),
                    f"{name}.txt", 'text/plain'
                )

        if st.button('Start New'):
            for k in list(st.session_state.keys()):
                if k.startswith('offer_'):
                    del st.session_state[k]
            st.rerun()


def admin_portal_ui(conn):
    """Enhanced admin portal with usage analytics, subscription management, and session management"""
    st.header("🔒 Admin Portal")

    tab1, tab2, tab3, tab4, tab5 = st.tabs(["User Management", "Subscription Management", "Content Management", "Usage Analytics", "Session Management"])

    with tab1:
        st.subheader("User Accounts")
        cursor = conn.cursor()
        cursor.execute("PRAGMA table_info(users)")
        columns = [column[1] for column in cursor.fetchall()]

        select_columns = ["username", "role"]
        if "last_login" in columns:
            select_columns.append("last_login")
        if "location_id" in columns:
            select_columns.append("location_id")
        if "created_at" in columns:
            select_columns.append("created_at")

        query = f"SELECT {', '.join(select_columns)} FROM users"
        users = conn.execute(query).fetchall()

        formatted_users = []
        for user in users:
            formatted_user = list(user)
            for i, col in enumerate(select_columns):
                if isinstance(formatted_user[i], str) and col in ['last_login', 'created_at']:
                    try:
                        formatted_user[i] = datetime.strptime(formatted_user[i], "%Y-%m-%d %H:%M:%S").strftime("%Y-%m-%d %H:%M")
                    except:
                        pass
            formatted_users.append(formatted_user)

        user_df = pd.DataFrame(formatted_users, columns=select_columns)
        st.dataframe(user_df)

        with st.expander("Create New User"):
            new_user = st.text_input("Username")
            new_pass = st.text_input("Password", type="password")
            user_role = st.selectbox("Role", ["user", "admin"])
            location_id = st.text_input("Location ID")

            if st.button("Add User"):
                if not new_user or not new_pass:
                    st.error("Username and password are required")
                elif len(new_pass) < 8:
                    st.error("Password must be at least 8 characters")
                else:
                    hashed = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt())
                    try:
                        if "location_id" in columns:
                            conn.execute(
                                "INSERT INTO users (username, password, role, location_id) VALUES (?, ?, ?, ?)",
                                (new_user, hashed, user_role, location_id)
                            )
                        else:
                            conn.execute(
                                "INSERT INTO users (username, password, role) VALUES (?, ?, ?)",
                                (new_user, hashed, user_role)
                            )
                        conn.commit()
                        st.success("User created successfully!")
                        time.sleep(1)
                        st.rerun()
                    except sqlite3.IntegrityError:
                        st.error("Username already exists")

        with st.expander("Delete User"):
            delete_user = st.selectbox("Select User to Delete", [u[0] for u in users if u[0] != "admin"])
            if st.button("Delete User"):
                if delete_user == st.session_state.username:
                    st.error("Cannot delete your own account!")
                else:
                    conn.execute("DELETE FROM users WHERE username = ?", (delete_user,))
                    delete_user_sessions(conn, delete_user)
                    conn.execute("DELETE FROM subscriptions WHERE username = ?", (delete_user,))
                    conn.execute("DELETE FROM interactions WHERE username = ?", (delete_user,))
                    conn.commit()
                    st.success(f"User '{delete_user}' deleted successfully!")
                    time.sleep(1)
                    st.rerun()

    with tab2:
        st.subheader("Feature Access Control")

        users = conn.execute("SELECT username FROM users").fetchall()
        if not users:
            st.warning("No users found")
        else:
            selected_user = st.selectbox("Select User", [u[0] for u in users])

            sub = conn.execute(
                "SELECT lease_analysis, deal_structuring, offer_generator FROM subscriptions WHERE username = ?",
                (selected_user,)
            ).fetchone()

            if not sub:
                conn.execute(
                    "INSERT INTO subscriptions (username) VALUES (?)",
                    (selected_user,)
                )
                conn.commit()
                sub = (0, 0, 0)

            col1, col2, col3 = st.columns(3)
            with col1:
                lease_access = st.toggle("Lease Analysis", value=bool(sub[0]))
            with col2:
                deal_access = st.toggle("Deal Structuring", value=bool(sub[1]))
            with col3:
                offer_access = st.toggle("Offer Generator", value=bool(sub[2]))

            if st.button("Update Access"):
                conn.execute(
                    """UPDATE subscriptions
                    SET lease_analysis = ?, deal_structuring = ?, offer_generator = ?
                    WHERE username = ?""",
                    (int(lease_access), int(deal_access), int(offer_access), selected_user)
                )
                conn.commit()
                st.success("Access updated successfully!")

    with tab3:
        st.subheader("Training Content")
        with st.expander("Upload Training Materials"):
            file_type = st.selectbox("Content Type", ["Document", "Video"])
            uploaded = st.file_uploader(
                f"Upload {file_type}",
                type=["pdf", "docx", "mp4"] if file_type == "Document" else ["mp4", "mov"]
            )
            description = st.text_area("Content Description")

            if uploaded and st.button("Upload"):
                save_dir = "training_content"
                os.makedirs(save_dir, exist_ok=True)
                file_path = os.path.join(save_dir, uploaded.name)

                with open(file_path, "wb") as f:
                    f.write(uploaded.getbuffer())

                meta_path = os.path.join(save_dir, f"{uploaded.name}.meta")
                with open(meta_path, "w") as f:
                    json.dump({
                        "uploaded_by": st.session_state.username,
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "description": description,
                        "type": file_type.lower()
                    }, f)

                st.success(f"{file_type} uploaded successfully!")

        st.subheader("Content Library")
        if os.path.exists("training_content"):
            files = os.listdir("training_content")
            content_files = [f for f in files if not f.endswith(".meta")]

            for file in content_files:
                meta_file = f"{file}.meta"
                if meta_file in files:
                    with open(os.path.join("training_content", meta_file)) as f:
                        meta = json.load(f)
                    st.markdown(f"**{file}**")
                    st.caption(f"Type: {meta['type']} | Uploaded by: {meta['uploaded_by']}")
                    st.caption(f"Description: {meta['description']}")
                    st.download_button(
                        f"Download {file}",
                        data=open(os.path.join("training_content", file), "rb").read(),
                        file_name=file
                    )
                    st.divider()

    with tab4:
        st.subheader("Usage Analytics")
        st.write("### Feature Usage")
        usage = conn.execute(
            "SELECT feature, COUNT(*) as count FROM interactions GROUP BY feature"
        ).fetchall()
        if usage:
            fig = px.pie(
                names=[u[0] for u in usage],
                values=[u[1] for u in usage],
                title="Feature Usage Distribution"
            )
            st.plotly_chart(fig)
        else:
            st.warning("No usage data available yet")

        st.write("### User Activity")
        activity = conn.execute(
            "SELECT username, COUNT(*) as interactions "
            "FROM interactions GROUP BY username ORDER BY interactions DESC LIMIT 10"
        ).fetchall()
        if activity:
            fig = px.bar(
                x=[a[0] for a in activity],
                y=[a[1] for a in activity],
                labels={"x": "User", "y": "Interactions"},
                title="Top Users by Activity"
            )
            st.plotly_chart(fig)
        else:
            st.warning("No user activity data available")

    with tab5:
        st.subheader("Active Sessions")
        sessions = conn.execute(
            "SELECT session_token, username, expires_at, created_at FROM sessions ORDER BY created_at DESC"
        ).fetchall()
        if sessions:
            session_data = [
                {
                    "Session Token": s[0][:8] + "...",
                    "Username": s[1],
                    "Expires At": datetime.fromisoformat(s[2]).strftime("%Y-%m-%d %H:%M"),
                    "Created At": datetime.fromisoformat(s[3]).strftime("%Y-%m-%d %H:%M")
                }
                for s in sessions
            ]
            st.dataframe(pd.DataFrame(session_data))
            selected_session = st.selectbox("Select Session to Terminate", [s[0] for s in sessions])
            if st.button("Terminate Session"):
                delete_session(conn, selected_session)
                st.success("Session terminated successfully!")
                st.rerun()
        else:
            st.info("No active sessions found.")


def history_ui(conn):
    """Show user's interaction history"""
    st.header("🕒 Your History")

    if "username" not in st.session_state:
        st.warning("Please log in to view your history")
        return

    if "current_interaction" in st.session_state:
        interaction = st.session_state.current_interaction
        st.subheader(f"Full Interaction – {interaction['timestamp']}")
        st.write(f"**Feature:** {interaction['feature']}")
        tabs = st.tabs(["Input", "Output"])
        with tabs[0]:
            st.text(interaction["input"])
        with tabs[1]:
            st.markdown(interaction["output"])

        if st.button("← Back to History"):
            del st.session_state.current_